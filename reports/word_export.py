# -*- coding: utf-8 -*-
"""
Модуль экспорта результатов сверки в формат Word (Акт сверки)
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from models import ReconciliationSummary

def set_cell_background(cell, color_hex):
    """
    Устанавливает фоновый цвет ячейки таблицы
    """
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def export_to_word(
    unmatched_tp: list,
    unmatched_bt: list,
    summary: ReconciliationSummary,
    output_path: str
) -> None:
    """
    Генерирует официальный Акт сверки взаиморасчетов в формате Word (.docx)
    """
    doc = Document()
    
    # Настройки полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Стиль текста по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Заголовок документа
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("АКТ СВЕРКИ ВЗАИМОРАСЧЕТОВ")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 73, 125)  # #1F497D
    
    # Подзаголовок
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"по состоянию на {datetime.now().strftime('%d.%m.%Y')}")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    
    doc.add_paragraph(
        "Настоящий Акт составлен программным комплексом «Умная сверка 3.0» на основе сопоставления реестра "
        "продаж компании ООО «ТИКЕТПРОФИ» и реестра приходов (себестоимости) компании ООО «БАРС ТУР»."
    )
    
    # 1. Раздел общей сводки
    doc.add_heading("1. Общие показатели взаиморасчетов", level=2)
    
    stats_table = doc.add_table(rows=5, cols=3)
    stats_table.style = 'Table Grid'
    
    headers = ["Показатель", "Количество операций", "Сумма (руб.)"]
    hdr_cells = stats_table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F497D")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                
    row_data = [
        ("Сведения по реестру ООО «ТИКЕТПРОФИ»", str(summary.total_tp_count), f"{summary.total_tp_sum:,.2f}"),
        ("Сведения по реестру ООО «БАРС ТУР»", str(summary.total_bt_count), f"{summary.total_bt_sum:,.2f}"),
        ("Сопоставленные (согласованные) позиции", str(summary.matched_tp_count), f"{summary.matched_tp_sum:,.2f}"),
        ("Выявленная расчетная прибыль (маржа)", "", f"{summary.total_profit:,.2f}")
    ]
    
    for r_idx, (p_name, count, total) in enumerate(row_data, 1):
        row_cells = stats_table.rows[r_idx].cells
        row_cells[0].text = p_name
        row_cells[1].text = count
        row_cells[2].text = total
        
        # Выравнивание
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Стили
        if r_idx == 4:
            for cell in row_cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        
    doc.add_paragraph()  # Отступ
    
    # 2. Раздел расхождений
    doc.add_heading("2. Выявленные несоответствия и расхождения", level=2)
    doc.add_paragraph(
        "Ниже приведен перечень позиций, по которым при сверке не были обнаружены соответствующие пары "
        "по сквозным идентификаторам или наименованиям услуг."
    )
    
    total_discrepancies = len(unmatched_tp) + len(unmatched_bt)
    
    if total_discrepancies == 0:
        doc.add_paragraph("Расхождений не обнаружено. Все позиции сопоставлены корректно.")
    else:
        disc_table = doc.add_table(rows=1, cols=5)
        disc_table.style = 'Table Grid'
        
        # Заголовок таблицы расхождений
        headers_disc = ["№", "Система происхождения", "Тип услуги", "Описание услуги", "Сумма (руб.)"]
        hdr_cells = disc_table.rows[0].cells
        for i, title in enumerate(headers_disc):
            hdr_cells[i].text = title
            set_cell_background(hdr_cells[i], "1F497D")
            for p in hdr_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    
        counter = 1
        
        # Только в TicketProf
        for tp in unmatched_tp:
            row_cells = disc_table.add_row().cells
            row_cells[0].text = str(counter)
            row_cells[1].text = "TicketProf"
            row_cells[2].text = tp.service_type
            row_cells[3].text = tp.desc
            row_cells[4].text = f"{tp.allocated_amount:,.2f}"
            
            # Фоновый цвет для расхождений
            for cell in row_cells:
                set_cell_background(cell, "FCE4D6")  # Оранжевый тон
                
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            counter += 1
            
        # Только в Bars Tour
        for bt in unmatched_bt:
            row_cells = disc_table.add_row().cells
            row_cells[0].text = str(counter)
            row_cells[1].text = "Bars Tour"
            row_cells[2].text = bt.service_type
            row_cells[3].text = bt.desc
            row_cells[4].text = f"{bt.amount:,.2f}"
            
            for cell in row_cells:
                set_cell_background(cell, "FCE4D6")
                
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            counter += 1
            
    doc.add_paragraph()  # Отступ
    
    # 3. Подписи сторон
    doc.add_paragraph("Акт составлен в двух экземплярах, имеющих одинаковую юридическую силу.")
    doc.add_paragraph()
    
    # Таблица для подписей
    sign_table = doc.add_table(rows=3, cols=2)
    sign_table.autofit = False
    
    # Задаем ширину колонок под подписи
    for row in sign_table.rows:
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(3.2)
        
    row_cells = sign_table.rows[0].cells
    row_cells[0].text = "От ООО «ТИКЕТПРОФИ»:"
    row_cells[1].text = "От ООО «БАРС ТУР»:"
    
    for i in range(2):
        row_cells[i].paragraphs[0].runs[0].font.bold = True
        
    sign_cells = sign_table.rows[2].cells
    sign_cells[0].text = "______________________ / ____________ /\n\nМ.П."
    sign_cells[1].text = "______________________ / ____________ /\n\nМ.П."
    
    doc.save(output_path)
